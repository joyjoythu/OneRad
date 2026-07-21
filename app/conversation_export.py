"""把会话消息渲染为 Markdown / Word 文档并写入项目目录。

输入为 _render_messages 产出的消息 dict 列表（role/content/timestamp/
tool_calls/tool_call_id/reasoning_content）。
"""
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

EXPORT_DIR_NAME = "conversation_exports"


def _now_local() -> datetime:
    return datetime.now(timezone.utc).astimezone()


def _fmt_time(iso: Optional[str]) -> str:
    """消息时间戳（UTC ISO）转本地可读时间；无效输入返回空串。"""
    if not iso:
        return ""
    try:
        return datetime.fromisoformat(iso).astimezone().strftime("%Y-%m-%d %H:%M:%S")
    except ValueError:
        return ""


def _tool_name_map(messages: List[Dict[str, Any]]) -> Dict[str, str]:
    """tool_call_id -> 工具名，用于给工具结果小节标注来源工具。"""
    names: Dict[str, str] = {}
    for msg in messages:
        for tc in msg.get("tool_calls") or []:
            if tc.get("id"):
                names[tc["id"]] = tc.get("name") or "unknown"
    return names


def _iter_sections(title: str, messages: List[Dict[str, Any]]):
    """把消息序列化为统一的小节流，供 Markdown 与 docx 两种渲染复用。

    每个小节为 dict：kind = meta / user / assistant / reasoning / tool_calls / tool。
    """
    yield {"kind": "meta", "text": f"导出时间：{_now_local().strftime('%Y-%m-%d %H:%M:%S')}"}
    names = _tool_name_map(messages)
    for msg in messages:
        role = msg.get("role")
        ts = _fmt_time(msg.get("timestamp"))
        suffix = f" · {ts}" if ts else ""
        if role == "user":
            yield {"kind": "user", "heading": f"用户{suffix}",
                   "text": msg.get("content") or ""}
        elif role == "assistant":
            yield {"kind": "assistant", "heading": f"助手{suffix}",
                   "text": msg.get("content") or ""}
            if msg.get("reasoning_content"):
                yield {"kind": "reasoning", "text": msg["reasoning_content"]}
            tool_calls = msg.get("tool_calls") or []
            if tool_calls:
                yield {"kind": "tool_calls", "items": [
                    {"name": tc.get("name") or "unknown",
                     "args": json.dumps(tc.get("args") or {}, ensure_ascii=False)}
                    for tc in tool_calls
                ]}
        elif role == "tool":
            name = names.get(msg.get("tool_call_id"), "工具")
            yield {"kind": "tool", "heading": f"工具结果：{name}{suffix}",
                   "text": msg.get("content") or ""}


def render_markdown(title: str, messages: List[Dict[str, Any]]) -> str:
    lines = [f"# 对话导出：{title}", ""]
    for section in _iter_sections(title, messages):
        kind = section["kind"]
        if kind == "meta":
            lines.append(f"> {section['text']}")
            lines.append("")
        elif kind in ("user", "assistant"):
            lines.append(f"## {section['heading']}")
            lines.append("")
            if section["text"]:
                lines.append(section["text"])
                lines.append("")
        elif kind == "reasoning":
            lines.append("**思考过程：**")
            lines.append("")
            for ln in section["text"].splitlines():
                lines.append(f"> {ln}")
            lines.append("")
        elif kind == "tool_calls":
            lines.append("**工具调用：**")
            lines.append("")
            for item in section["items"]:
                lines.append(f"- `{item['name']}` `{item['args']}`")
            lines.append("")
        elif kind == "tool":
            lines.append(f"### {section['heading']}")
            lines.append("")
            lines.append("```")
            lines.append(section["text"])
            lines.append("```")
            lines.append("")
    return "\n".join(lines)


def write_docx(title: str, messages: List[Dict[str, Any]], path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading(f"对话导出：{title}", level=0)
    for section in _iter_sections(title, messages):
        kind = section["kind"]
        if kind == "meta":
            doc.add_paragraph(section["text"])
        elif kind in ("user", "assistant"):
            doc.add_heading(section["heading"], level=2)
            if section["text"]:
                doc.add_paragraph(section["text"])
        elif kind == "reasoning":
            para = doc.add_paragraph()
            para.add_run("思考过程：").bold = True
            run = doc.add_paragraph().add_run(section["text"])
            run.italic = True
        elif kind == "tool_calls":
            para = doc.add_paragraph()
            para.add_run("工具调用：").bold = True
            for item in section["items"]:
                run = doc.add_paragraph(style="List Bullet").add_run(
                    f"{item['name']} {item['args']}"
                )
                run.font.name = "Consolas"
        elif kind == "tool":
            doc.add_heading(section["heading"], level=3)
            run = doc.add_paragraph().add_run(section["text"])
            run.font.name = "Consolas"
    doc.save(str(path))


def _safe_filename(text: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|]', "_", text).strip()
    return cleaned or "conversation"


def export_conversation(
    project_path: str,
    title: str,
    messages: List[Dict[str, Any]],
    fmt: str,
) -> Path:
    """把会话导出为 md/docx 文件，写入项目目录下 conversation_exports/。"""
    out_dir = Path(project_path) / EXPORT_DIR_NAME
    out_dir.mkdir(parents=True, exist_ok=True)
    stamp = _now_local().strftime("%Y%m%d_%H%M%S")
    stem = f"对话_{_safe_filename(title)}_{stamp}"
    if fmt == "docx":
        path = out_dir / f"{stem}.docx"
        write_docx(title, messages, path)
    else:
        path = out_dir / f"{stem}.md"
        path.write_text(render_markdown(title, messages), encoding="utf-8")
    return path
