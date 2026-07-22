"""中文学术论文格式的 docx 排版模块。

集中管理 Word 报告的格式规范：正文宋体小四、西文/数字 Times New Roman、
1.5 倍行距、标题黑体、表格五号。提供三个入口：

- ``apply_academic_style(doc)``：在新建文档上设置样式级格式；
- ``reformat_docx(path)``：原地重排已有 docx（不生成备份），幂等；
- ``markdown_to_docx(doc, text)``：把 markdown 片段渲染为 Word 段落。

技术要点：中文字体必须同时在 ``style.font.name``（西文）之外，通过
``rPr.rFonts`` 的 ``w:eastAsia`` 属性设置；python-docx 默认主题的 Heading
样式是蓝色，需要显式覆盖为黑色。
"""

import re

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SONG_TI = "宋体"
HEI_TI = "黑体"
TIMES_NEW_ROMAN = "Times New Roman"
BLACK = RGBColor(0, 0, 0)

# 标题样式规范：中文字体、字号（pt）、加粗；西文/数字统一 Times New Roman。
_HEADING_SPECS = {
    "Title": {"east": HEI_TI, "size": 16, "bold": True},
    "Heading 1": {"east": HEI_TI, "size": 14, "bold": True},
    "Heading 2": {"east": HEI_TI, "size": 12, "bold": True},
}
# 正文（Normal / List Bullet 等）：宋体小四 12pt。
_BODY_SPEC = {"east": SONG_TI, "size": 12, "bold": False}
BODY_LINE_SPACING = 1.5
TABLE_FONT_SIZE = 10.5  # 五号


def _set_east_asia(rpr_holder, east_font: str) -> None:
    """在 run/style 元素上设置中文字体（w:eastAsia）。"""
    rPr = rpr_holder.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), east_font)


def _apply_style_font(style, east_font: str, size: float, bold: bool) -> None:
    """在样式层面统一设置西文字体、中文字体、字号、加粗与黑色。"""
    style.font.name = TIMES_NEW_ROMAN
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = BLACK  # 显式黑色：覆盖 Heading 默认的主题蓝
    _set_east_asia(style.element, east_font)


def _apply_run_font(run, east_font: str, size: float,
                    bold=None) -> None:
    """在 run 层面归一化字体/字号/颜色；bold=None 时保留原加粗设置。"""
    run.font.name = TIMES_NEW_ROMAN
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = BLACK
    _set_east_asia(run._element, east_font)


def apply_academic_style(doc) -> None:
    """为新建文档设置学术论文样式（样式级，作用于随后写入的所有内容）。

    Normal / List Bullet：宋体 + Times New Roman、小四 12pt、黑色、1.5 倍行距；
    Title：黑体三号 16pt 加粗居中；Heading 1：黑体四号 14pt 加粗黑色；
    Heading 2：黑体小四 12pt 加粗。
    """
    for name in ("Normal", "List Bullet"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        _apply_style_font(style, _BODY_SPEC["east"], _BODY_SPEC["size"],
                          _BODY_SPEC["bold"])
        style.paragraph_format.line_spacing = BODY_LINE_SPACING

    for name, spec in _HEADING_SPECS.items():
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        _apply_style_font(style, spec["east"], spec["size"], spec["bold"])
    try:
        doc.styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except KeyError:
        pass


def style_table(table) -> None:
    """统一表格字体：全部单元格五号 10.5pt，表头加粗居中。

    表格内容字体不随 Normal 样式变化，必须在单元格段落的 run 上设置。
    """
    for row_idx, row in enumerate(table.rows):
        is_header = row_idx == 0
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if is_header:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _apply_run_font(run, SONG_TI, TABLE_FONT_SIZE,
                                    bold=True if is_header else None)


def _add_markdown_runs(paragraph, text: str) -> None:
    """把含 ``**粗体**`` 标记的文本写入段落，粗体片段保留加粗。"""
    for idx, part in enumerate(text.split("**")):
        if not part:
            continue
        run = paragraph.add_run(part)
        if idx % 2 == 1:
            run.bold = True


_HEADING_STYLES = {1: "Title", 2: "Heading 1", 3: "Heading 2"}


def markdown_to_docx(doc, text: str) -> None:
    """把 markdown 片段渲染为 Word 内容。

    ``#``/``##``/``###`` → Title/Heading 1/Heading 2，``- ``/``* `` →
    List Bullet，``**粗体**`` 解析为加粗 run（不丢弃），普通行 → 正文段落，
    空行跳过。
    """
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            paragraph = doc.add_paragraph(style=_HEADING_STYLES[level])
            _add_markdown_runs(paragraph, heading.group(2).strip())
        elif line.startswith(("- ", "* ")):
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_markdown_runs(paragraph, line[2:].strip())
        else:
            paragraph = doc.add_paragraph()
            _add_markdown_runs(paragraph, line)


def reformat_docx(path: str) -> str:
    """原地重排已有 docx 为学术论文格式，返回该文件路径。

    按段落 style 名（Title/Heading 1/Heading 2/Normal/List Bullet 等）
    归一化所有 run 的字体/字号/颜色（覆盖显式 Pt() 等 run 级设置；正文 run
    的加粗保持不变），表格单元格统一五号。直接保存到原文件，不生成
    .bak 备份。重复执行幂等。
    """
    from docx import Document
    doc = Document(path)
    apply_academic_style(doc)
    for paragraph in doc.paragraphs:
        spec = _HEADING_SPECS.get(paragraph.style.name)
        if spec is not None:
            for run in paragraph.runs:
                _apply_run_font(run, spec["east"], spec["size"],
                                bold=spec["bold"])
            if paragraph.style.name == "Title":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            for run in paragraph.runs:
                _apply_run_font(run, _BODY_SPEC["east"], _BODY_SPEC["size"])
    for table in doc.tables:
        style_table(table)
    doc.save(path)
    return path
