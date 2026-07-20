"""Feature statistics analysis for the conversational agent.

Provides statistical comparison of selected radiomic features between
two label groups (0/1): independent-samples t-test and Mann-Whitney U test,
with descriptive statistics and a Word table output.

Two entry points:

- ``inspect_statistics_inputs``: resolve feature CSV / clinical table /
  selected-features CSV / ID and label columns. Returns ``ready`` with
  resolved parameters, or ``need_clarification`` / ``error`` so the agent
  can ask the user before any execution.
- ``run_feature_statistics``: run the statistical analysis and export
  a Word table + CSV.
"""

import logging
import os
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from scipy import stats

from app.radiomics_analysis import _load_table
from app.utils import _load_feature_csv, _merge_feature_clinical

logger = logging.getLogger(__name__)

_RADIOMIC_PREFIXES = ("original_", "wavelet-", "log-sigma_")
_DEFAULT_OUTPUT_DIR = "feature_statistics"


def inspect_statistics_inputs(
    project_path: str,
    feature_csv: str = "",
    clinical: str = "",
    id_col: str = "",
    label_col: str = "",
    selected_features_csv: str = "",
    output_dir: str = "",
) -> Dict[str, Any]:
    """Resolve statistics inputs; report ambiguity as clarification questions.

    Returns one of:
      - {"status": "ready", "resolved": {...}}  all inputs resolved
      - {"status": "need_clarification", "questions": [...], "detected": {...}}
      - {"status": "error", "message": str, "detected": {...}}
    """
    # 1. 特征文件
    if not feature_csv:
        feature_csv = os.path.join(
            project_path, "radiomics_features", "radiomics_features.csv")
    try:
        feature_df = _load_feature_csv(feature_csv)
    except (FileNotFoundError, ValueError) as e:
        return {"status": "error",
                "message": f"{e}。请先提取特征或通过 feature_csv 指定路径",
                "detected": {}}
    feature_ids = set(feature_df["patient_id"].astype(str).str.strip())
    n_features = len([c for c in feature_df.columns
                      if any(c.startswith(p) for p in _RADIOMIC_PREFIXES)])
    detected: Dict[str, Any] = {
        "feature_csv": feature_csv,
        "n_feature_cases": int(len(feature_df)),
        "n_features": n_features,
    }

    # 2. 临床文件 — 复用 radiomics_analysis 的扫描逻辑
    from app.radiomics_analysis import (
        _find_clinical_candidates,
        _binary_columns,
        _detect_id_column,
        _norm_id,
    )

    if not clinical:
        candidates = _find_clinical_candidates(
            project_path, feature_ids, exclude_path=feature_csv)
        if not candidates:
            return {"status": "error",
                    "message": "未在项目内找到可用的临床表格（需含 0/1 标签列且 "
                               "ID 与特征匹配），请通过 clinical 指定路径",
                    "detected": detected}
        if len(candidates) > 1:
            rel = [os.path.relpath(c, project_path) for c in candidates]
            return {"status": "need_clarification",
                    "questions": [{"field": "clinical",
                                   "question": "找到多个可能的临床表格，请指定使用哪一个",
                                   "candidates": rel}],
                    "detected": detected}
        clinical = candidates[0]
    try:
        clinical_df = _load_table(clinical)
    except (FileNotFoundError, ValueError) as e:
        return {"status": "error", "message": str(e), "detected": detected}
    except Exception as e:
        return {"status": "error",
                "message": f"读取临床表格失败: {e}",
                "detected": detected}
    detected["clinical"] = clinical

    questions: List[Dict[str, Any]] = []

    # 3. ID 列
    from app.radiomics_analysis import _id_match_counts
    counts = _id_match_counts(clinical_df, feature_ids)
    if id_col:
        if id_col not in clinical_df.columns:
            return {"status": "error",
                    "message": f"指定的 ID 列 '{id_col}' 不存在",
                    "detected": detected}
        if counts.get(id_col, 0) == 0:
            return {"status": "error",
                    "message": f"指定的 ID 列 '{id_col}' 与特征 patient_id 无任何匹配",
                    "detected": detected}
    else:
        best = max(counts, key=counts.get) if counts else None
        best_n = counts.get(best, 0) if best else 0
        if best is None or best_n == 0:
            questions.append({
                "field": "id_col",
                "question": "临床表中没有任何列能与特征 patient_id 匹配，请指定 ID 列",
                "candidates": list(clinical_df.columns),
            })
        else:
            ties = sorted([c for c, n in counts.items() if n == best_n])
            if len(ties) > 1:
                questions.append({
                    "field": "id_col",
                    "question": f"多列与特征 patient_id 的匹配数相同（{best_n} 例），请指定 ID 列",
                    "candidates": ties,
                })
            else:
                id_col = best
    if id_col:
        detected["id_col"] = id_col
        detected["n_matched"] = counts.get(id_col, 0)

    # 4. 标签列
    from app.radiomics_analysis import _binary_values, _binary_columns
    if label_col:
        if label_col not in clinical_df.columns:
            return {"status": "error",
                    "message": f"指定的标签列 '{label_col}' 不存在",
                    "detected": detected}
        values = _binary_values(clinical_df[label_col])
        if values is None:
            return {"status": "error",
                    "message": f"标签列 '{label_col}' 必须为 0/1 二分类",
                    "detected": detected}
        if len(values) < 2:
            return {"status": "error",
                    "message": f"标签列 '{label_col}' 必须同时包含 0 和 1",
                    "detected": detected}
    else:
        binary_cols = _binary_columns(clinical_df, exclude=id_col or None)
        named = [c for c in binary_cols if c.lower() == "label"]
        if named:
            label_col = named[0]
        elif len(binary_cols) == 1:
            label_col = binary_cols[0]
        elif len(binary_cols) > 1:
            questions.append({
                "field": "label_col",
                "question": "临床表中有多个 0/1 列，哪一列是分组标签？",
                "candidates": binary_cols,
            })
        else:
            questions.append({
                "field": "label_col",
                "question": "未找到 0/1 二分类标签列，请指定 label_col",
                "candidates": list(clinical_df.columns),
            })
    if label_col:
        detected["label_col"] = label_col

    # 5. 筛选特征文件
    if not selected_features_csv:
        # 尝试默认路径
        default = os.path.join(project_path, "radiomics_analysis", "selected_features.csv")
        if os.path.exists(default):
            selected_features_csv = default
        else:
            questions.append({
                "field": "selected_features_csv",
                "question": "未找到筛选特征文件 (selected_features.csv)，请指定路径",
            })
    if selected_features_csv:
        try:
            sel_df = pd.read_csv(selected_features_csv)
            feat_col = "feature" if "feature" in sel_df.columns else sel_df.columns[0]
            detected["selected_features"] = sel_df[feat_col].dropna().astype(str).tolist()
            detected["n_selected"] = len(detected["selected_features"])
            detected["selected_features_csv"] = selected_features_csv
        except Exception as e:
            return {"status": "error",
                    "message": f"读取筛选特征文件失败: {e}",
                    "detected": detected}

    if questions:
        return {"status": "need_clarification",
                "questions": questions, "detected": detected}

    # 6. 输出目录
    if not output_dir:
        output_dir = os.path.join(project_path, _DEFAULT_OUTPUT_DIR)

    return {"status": "ready", "resolved": {
        "feature_csv": feature_csv,
        "clinical": clinical,
        "id_col": id_col,
        "label_col": label_col,
        "selected_features_csv": selected_features_csv,
        "selected_features": detected.get("selected_features", []),
        "output_dir": output_dir,
        "n_feature_cases": detected["n_feature_cases"],
        "n_matched": counts.get(id_col, 0),
        "n_selected": detected.get("n_selected", 0),
    }}


def run_feature_statistics(
    feature_csv: str,
    clinical: str,
    id_col: str,
    label_col: str,
    selected_features: List[str],
    output_dir: str,
) -> Dict[str, Any]:
    """Run statistical comparison of selected features between label groups.

    Args:
        feature_csv: Path to radiomics features CSV.
        clinical: Path to clinical table.
        id_col: Name of the patient ID column in the clinical table.
        label_col: Name of the binary label column (0/1).
        selected_features: List of feature names to analyze.
        output_dir: Directory for output files.

    Returns:
        {"success": bool, "message": str, "outputs": {"csv": ..., "docx": ...},
         "n_features_analyzed": int, "n_significant_ttest": int,
         "n_significant_mwu": int, "results": [...]}
    """
    try:
        feature_df = _load_feature_csv(feature_csv)
    except (FileNotFoundError, ValueError) as e:
        return {"success": False, "message": str(e)}

    try:
        clinical_df = _load_table(clinical)
    except Exception as e:
        return {"success": False, "message": f"读取临床表格失败: {e}"}

    if id_col != "patient_id":
        clinical_df = clinical_df.rename(columns={id_col: "patient_id"})

    try:
        merged = _merge_feature_clinical(feature_df, clinical_df, id_col="patient_id")
    except ValueError as e:
        return {"success": False, "message": str(e)}

    if label_col not in merged.columns:
        return {"success": False, "message": f"标签列 '{label_col}' 不存在于合并后数据"}
    if merged[label_col].nunique() < 2:
        return {"success": False, "message": f"标签列 '{label_col}' 必须同时包含 0 和 1"}

    y = merged[label_col].astype(int)
    idx_0 = y == 0
    idx_1 = y == 1

    if idx_0.sum() == 0 or idx_1.sum() == 0:
        return {"success": False,
                "message": f"某组样本数为 0（0: {idx_0.sum()}, 1: {idx_1.sum()}）"}

    # Filter to features that exist in the merged data
    available = set(merged.columns)
    features_to_test = [f for f in selected_features if f in available]
    missing = [f for f in selected_features if f not in available]
    if missing:
        logger.warning("以下筛选特征不在特征矩阵中，已跳过: %s", missing)

    if not features_to_test:
        return {"success": False, "message": "筛选特征列表中没有任何特征存在于特征矩阵"}

    results = []
    sig_ttest = 0
    sig_mwu = 0

    for feat in features_to_test:
        g0 = merged.loc[idx_0, feat].dropna().values.astype(float)
        g1 = merged.loc[idx_1, feat].dropna().values.astype(float)

        if len(g0) < 2 or len(g1) < 2:
            # 任一组样本不足 2 例时跳过检验
            results.append({
                "feature": feat,
                "group0_n": len(g0), "group0_mean": float(np.mean(g0)) if len(g0) else None,
                "group0_std": float(np.std(g0, ddof=1)) if len(g0) > 1 else None,
                "group1_n": len(g1), "group1_mean": float(np.mean(g1)) if len(g1) else None,
                "group1_std": float(np.std(g1, ddof=1)) if len(g1) > 1 else None,
                "t_stat": None, "t_pvalue": None,
                "mw_u": None, "mw_pvalue": None,
                "note": "样本量不足",
            })
            continue

        g0_mean, g0_std = float(np.mean(g0)), float(np.std(g0, ddof=1))
        g1_mean, g1_std = float(np.mean(g1)), float(np.std(g1, ddof=1))

        # Independent-samples t-test
        try:
            t_stat, t_p = stats.ttest_ind(g0, g1, equal_var=False)  # Welch's t-test
            t_stat, t_p = float(t_stat), float(t_p)
        except Exception:
            t_stat, t_p = None, None

        # Mann-Whitney U test
        try:
            mw_u, mw_p = stats.mannwhitneyu(g0, g1, alternative="two-sided")
            mw_u, mw_p = float(mw_u), float(mw_p)
        except Exception:
            mw_u, mw_p = None, None

        if t_p is not None and t_p < 0.05:
            sig_ttest += 1
        if mw_p is not None and mw_p < 0.05:
            sig_mwu += 1

        results.append({
            "feature": feat,
            "group0_n": len(g0), "group0_mean": g0_mean, "group0_std": g0_std,
            "group1_n": len(g1), "group1_mean": g1_mean, "group1_std": g1_std,
            "t_stat": t_stat, "t_pvalue": t_p,
            "mw_u": mw_u, "mw_pvalue": mw_p,
        })

    # Export CSV
    os.makedirs(output_dir, exist_ok=True)
    csv_path = os.path.join(output_dir, "feature_statistics.csv")
    csv_df = pd.DataFrame(results)
    csv_df.to_csv(csv_path, index=False, encoding="utf-8-sig")

    # Export Word table
    docx_path = os.path.join(output_dir, "feature_statistics_report.docx")
    _write_word_report(results, idx_0.sum(), idx_1.sum(), docx_path)

    return {
        "success": True,
        "message": f"统计分析完成：共分析 {len(results)} 个特征，"
                   f"t检验显著 {sig_ttest} 个，Mann-Whitney U 显著 {sig_mwu} 个",
        "n_features_analyzed": len(results),
        "n_significant_ttest": sig_ttest,
        "n_significant_mwu": sig_mwu,
        "n_missing_features": len(missing),
        "outputs": {"csv": csv_path, "docx": docx_path},
        "results": results,
    }


def _write_word_report(
    results: List[Dict[str, Any]],
    n_group0: int,
    n_group1: int,
    output_path: str,
) -> None:
    """Write a Word document with a feature statistics summary table."""
    doc = Document()

    # Title
    title = doc.add_heading(level=0)
    run = title.add_run("影像组学特征统计分析报告")
    run.font.size = Pt(18)
    run.bold = True

    # Summary paragraph
    doc.add_paragraph(
        f"对 {len(results)} 个筛选后影像组学特征进行分组统计分析。"
        f"Label=0: {n_group0} 例，Label=1: {n_group1} 例。"
        f"采用独立样本 t 检验（Welch 校正，不假设方差齐性）"
        f"与 Mann-Whitney U 检验（非参数秩和检验）。"
    )

    sig_ttest = sum(1 for r in results if r.get("t_pvalue") is not None and r["t_pvalue"] < 0.05)
    sig_mwu = sum(1 for r in results if r.get("mw_pvalue") is not None and r["mw_pvalue"] < 0.05)
    doc.add_paragraph(
        f"t 检验显著（p<0.05）的特征：{sig_ttest}/{len(results)}；"
        f"Mann-Whitney U 检验显著（p<0.05）的特征：{sig_mwu}/{len(results)}。"
    )

    # Feature statistics table
    doc.add_heading("特征分组统计与检验结果", level=1)

    columns = [
        "特征名",
        "Group0 Mean±SD",
        "Group1 Mean±SD",
        "t值",
        "p值 (t检验)",
        "U值",
        "p值 (MW-U)",
    ]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr[i].text = col
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r in results:
        cells = table.add_row().cells
        cells[0].text = r["feature"]

        if r.get("group0_mean") is not None and r.get("group0_std") is not None:
            cells[1].text = f"{r['group0_mean']:.4f} ± {r['group0_std']:.4f}"
        else:
            cells[1].text = "N/A"

        if r.get("group1_mean") is not None and r.get("group1_std") is not None:
            cells[2].text = f"{r['group1_mean']:.4f} ± {r['group1_std']:.4f}"
        else:
            cells[2].text = "N/A"

        cells[3].text = f"{r['t_stat']:.4f}" if r.get("t_stat") is not None else "-"
        cells[4].text = f"{r['t_pvalue']:.4f}" if r.get("t_pvalue") is not None else "-"
        cells[5].text = f"{r['mw_u']:.1f}" if r.get("mw_u") is not None else "-"
        cells[6].text = f"{r['mw_pvalue']:.4f}" if r.get("mw_pvalue") is not None else "-"

        # Apply font size
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.save(output_path)
    logger.info("统计分析报告已保存: %s", output_path)
