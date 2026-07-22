import os
from typing import Dict, Any, List, Optional

import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import logging
from app.docx_style import apply_academic_style, markdown_to_docx, style_table
from app.skills import load_skill
from app.utils import fmt_num, fmt_p

logger = logging.getLogger(__name__)

# docx 与 report.md 共用的 SHAP 小节说明（口径一致）
SHAP_EXPLANATION = (
    "SHAP（SHapley Additive exPlanations）逐折量化每个特征对每例患者预测结果的贡献。"
    "蜂群图中每个点代表一例患者：点的颜色表示特征取值（红色为高、蓝色为低），"
    "横向位置表示该特征将预测推向阳性的方向和幅度；"
    "条形图按平均绝对 SHAP 值（贡献的平均幅度）对特征排序。"
)


class ReportAgent:
    """Generate a Word report from a radiomics analysis result."""

    def run(self, analysis_result: Dict[str, Any], output_dir: str,
            modality: str, n_features: int, covariates: List[str],
            plot_paths: Optional[List[str]] = None,
            llm_client=None,
            interpretation: Optional[Dict[str, str]] = None,
            covariate_display: Optional[List[str]] = None) -> Dict[str, Any]:
        """Build and save the DOCX report.

        Parameters
        ----------
        analysis_result: dict
            Output from the analysis agent. Must contain ``success=True`` and
            the required fields ``n_samples``, ``selected_features``,
            ``model_results`` (with ``coefficients``, ``odds_ratios``,
            ``ci_lower``, ``ci_upper``, ``p_values``) and ``metrics`` (with
            ``auc``, ``auc_ci``, ``accuracy``, ``sensitivity``,
            ``specificity``).
        output_dir: str
            Directory where the report will be saved.
        modality: str
            Imaging modality (e.g. ``"CT"``).
        n_features: int
            Total number of radiomic features extracted.
        covariates: list[str]
            Clinical covariates included in the logistic regression.
        plot_paths: list[str] | None
            Optional paths to figures to embed. Non-existent paths are skipped
            and reported in the returned ``skipped_plots`` list.
        llm_client: object | None
            Optional client with a ``call(system, prompt, ...)`` method used to
            polish the methodology paragraph.
        interpretation: dict | None
            Optional LLM-generated interpretation with ``performance`` /
            ``features`` / ``shap`` markdown strings. When provided, a
            ``结果解读`` section with three subsections is appended; ``None``
            keeps the current behaviour (no interpretation section).
        covariate_display: list[str] | None
            Optional display names for the covariates (e.g. translated
            ``Age（年龄）`` pairs); defaults to ``covariates`` when None.

        Returns
        -------
        dict
            ``{"success": bool, "message": str, "report_path": str,
            "skipped_plots": list[str]}``.
        """
        if analysis_result.get("success") is not True:
            return {
                "success": False,
                "message": f"上游分析失败: {analysis_result.get('message', '未知错误')}",
            }

        missing = []
        for key in ("n_samples", "selected_features", "model_results", "metrics"):
            if key not in analysis_result:
                missing.append(key)
        if "model_results" in analysis_result:
            for key in ("coefficients", "odds_ratios", "ci_lower", "ci_upper", "p_values"):
                if key not in analysis_result["model_results"]:
                    missing.append(f"model_results.{key}")
        if "metrics" in analysis_result:
            for key in ("auc", "auc_ci", "accuracy", "sensitivity", "specificity"):
                if key not in analysis_result["metrics"]:
                    missing.append(f"metrics.{key}")
        if missing:
            return {
                "success": False,
                "message": f"analysis_result 缺少字段: {', '.join(missing)}",
            }

        plot_paths = plot_paths or []

        try:
            os.makedirs(output_dir, exist_ok=True)
            doc = Document()
            apply_academic_style(doc)
            skipped_plots: List[str] = []

            # Title（字号/加粗/居中由 Title 样式统一控制）
            title = doc.add_heading(level=0)
            title.add_run("影像组学分析报告")

            # Methodology
            doc.add_heading("1. 方法", level=1)
            methodology = self._build_methodology(
                analysis_result["n_samples"], modality, n_features,
                len(analysis_result["selected_features"]), covariates,
                analysis_result["metrics"]["auc"],
                covariate_display=covariate_display,
            )
            if llm_client:
                methodology = self._polish_methodology(methodology, llm_client)
            doc.add_paragraph(methodology)

            # Feature Selection Table
            doc.add_heading("2. 特征选择", level=1)
            feat_df = pd.DataFrame({
                "特征名": analysis_result["selected_features"],
                "系数": [fmt_num(analysis_result["model_results"]["coefficients"].get(f, 0)) for f in analysis_result["selected_features"]],
            })
            self._add_table(doc, feat_df)

            # Regression Table
            doc.add_heading("3. 回归结果", level=1)
            rows = []
            for feat in analysis_result["selected_features"]:
                rows.append({
                    "特征": feat,
                    "OR": fmt_num(analysis_result['model_results']['odds_ratios'].get(feat, 0)),
                    "95%CI 下限": fmt_num(analysis_result['model_results']['ci_lower'].get(feat)),
                    "95%CI 上限": fmt_num(analysis_result['model_results']['ci_upper'].get(feat)),
                    "p 值": fmt_p(analysis_result['model_results']['p_values'].get(feat, 1)),
                })
            self._add_table(doc, pd.DataFrame(rows))

            # Performance
            doc.add_heading("4. 模型性能", level=1)
            m = analysis_result["metrics"]
            perf_text = (
                f"逻辑回归模型的 AUC 为 {m['auc']:.3f}"
                f"（95% CI：{m['auc_ci'][0]:.3f}–{m['auc_ci'][1]:.3f}）。"
                f"准确率 = {m['accuracy']:.3f}，敏感度 = {m['sensitivity']:.3f}，"
                f"特异度 = {m['specificity']:.3f}，"
                f"PPV = {fmt_num(m.get('ppv'))}，NPV = {fmt_num(m.get('npv'))}，"
                f"F1 = {fmt_num(m.get('f1'))}。"
            )
            doc.add_paragraph(perf_text)

            # Per-fold cross-validation metrics
            cv = analysis_result.get("cv_metrics")
            if cv and cv.get("folds"):
                doc.add_heading("5. 交叉验证详情", level=1)
                fold_rows = []
                for f in cv["folds"]:
                    fold_rows.append({
                        "折": f["fold"],
                        "AUC": fmt_num(f["auc"]),
                        "准确率": fmt_num(f["accuracy"]),
                        "敏感度": fmt_num(f["sensitivity"]),
                        "特异度": fmt_num(f["specificity"]),
                        "PPV": fmt_num(f["ppv"]),
                        "NPV": fmt_num(f["npv"]),
                        "F1": fmt_num(f["f1"]),
                    })
                mean, std = cv["mean"], cv["std"]
                fold_rows.append({
                    "折": "均值±标准差",
                    "AUC": f"{fmt_num(mean['auc'])}±{fmt_num(std['auc'])}",
                    "准确率": f"{fmt_num(mean['accuracy'])}±{fmt_num(std['accuracy'])}",
                    "敏感度": f"{fmt_num(mean['sensitivity'])}±{fmt_num(std['sensitivity'])}",
                    "特异度": f"{fmt_num(mean['specificity'])}±{fmt_num(std['specificity'])}",
                    "PPV": f"{fmt_num(mean['ppv'])}±{fmt_num(std['ppv'])}",
                    "NPV": f"{fmt_num(mean['npv'])}±{fmt_num(std['npv'])}",
                    "F1": f"{fmt_num(mean['f1'])}±{fmt_num(std['f1'])}",
                })
                self._add_table(doc, pd.DataFrame(fold_rows))

            # Visualizations（SHAP 图单独成节；LASSO path 仅展示 fold1 一张）
            shap_plots = [p for p in plot_paths
                          if os.path.basename(p).startswith(
                              ("shap_summary_fold", "shap_bar_fold"))]
            other_plots = [p for p in plot_paths if p not in shap_plots]
            if other_plots:
                doc.add_heading("6. 图表", level=1)
                if any(os.path.basename(p).startswith("lasso_path_fold")
                       for p in other_plots):
                    doc.add_paragraph(
                        "仅展示 fold 1 的 LASSO path 作为代表；"
                        "所有折的 LASSO path 均已保存在 lasso/ 目录中。")
                for plot_path in other_plots:
                    if os.path.exists(plot_path):
                        doc.add_picture(plot_path, width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        skipped_plots.append(plot_path)

            # SHAP Interpretability
            if shap_plots:
                doc.add_heading("7. SHAP 可解释性", level=1)
                doc.add_paragraph(SHAP_EXPLANATION)
                for plot_path in shap_plots:
                    if os.path.exists(plot_path):
                        doc.add_picture(plot_path, width=Inches(5.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        skipped_plots.append(plot_path)

            # 结果解读（LLM 生成的中文解读；interpretation 为 None 时不生成该节，
            # 报告整体重写，重复生成幂等）
            if interpretation:
                doc.add_heading("结果解读", level=1)
                for subtitle, key in (("模型性能解读", "performance"),
                                      ("特征意义解读", "features"),
                                      ("SHAP 可解释性解读", "shap")):
                    text = (interpretation.get(key) or "").strip()
                    if not text:
                        continue
                    doc.add_heading(subtitle, level=2)
                    markdown_to_docx(doc, text)

            # Save
            report_path = os.path.join(output_dir, "AutoRadiomics_Report.docx")
            doc.save(report_path)
            return {
                "success": True,
                "message": "报告生成完成",
                "report_path": report_path,
                "skipped_plots": skipped_plots,
            }
        except Exception as e:
            return {"success": False, "message": f"报告生成失败: {e}"}

    def _build_methodology(
        self,
        n_samples: int,
        modality: str,
        n_features: int,
        n_selected: int,
        covariates: List[str],
        auc: float,
        covariate_display: Optional[List[str]] = None,
    ) -> str:
        """Compose the methodology paragraph for the report.

        Parameters
        ----------
        n_samples: int
            Number of patients included.
        modality: str
            Imaging modality.
        n_features: int
            Total number of extracted radiomic features.
        n_selected: int
            Number of features selected by LASSO.
        covariates: list[str]
            Clinical covariates; omitted from the text when empty.
        auc: float
            Model AUC value.
        covariate_display: list[str] | None
            Optional display names for the covariates (e.g. translated
            ``Age（年龄）`` pairs); defaults to ``covariates`` when None.

        Returns
        -------
        str
            Methodology paragraph.
        """
        display = covariate_display if covariate_display is not None else covariates
        cov_str = ", ".join(display)
        cov_clause = (
            f"，并纳入临床协变量 {cov_str}" if covariates else ""
        )
        return (
            f"共纳入 {n_samples} 例患者。"
            f"使用 PyRadiomics 从 {modality} 图像中提取影像组学特征，共 {n_features} 个；"
            f"经 LASSO 回归筛选出 {n_selected} 个特征，进入逻辑回归模型{cov_clause}。"
            f"模型 AUC 为 {auc:.3f}。"
        )

    def _polish_methodology(self, raw: str, llm_client) -> str:
        """Polish the methodology paragraph using an optional LLM client.

        Parameters
        ----------
        raw: str
            Original methodology text.
        llm_client: object
            Client exposing a ``call(system, prompt, ...)`` method.

        Returns
        -------
        str
            Polished text, or ``raw`` if the call fails or returns an empty
            response.
        """
        system = load_skill("report-writing")
        try:
            polished = llm_client.call(system, raw, temperature=0.3, max_tokens=800)
            return polished or raw
        except Exception:
            return raw

    def _add_table(self, doc, df: pd.DataFrame):
        """Append a pandas DataFrame as a styled Word table.

        Parameters
        ----------
        doc: docx.Document
            Document to append the table to.
        df: pandas.DataFrame
            Data to render. Column names become the table header.
        """
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col)
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        style_table(table)
