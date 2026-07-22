"""面向 agent 工具的 Word 文档操作入口：创建、追加、重排。

路径安全（限制在项目目录内）由调用方（agent 工具层）负责；本模块只做
文件级操作，统一返回 ``{"success": bool, ...}`` 字典。
"""

import os

from docx import Document

from app.docx_style import apply_academic_style, markdown_to_docx, reformat_docx


def create_document(path: str, markdown: str, title: str = None) -> dict:
    """新建套用学术格式的 docx；文件已存在时报错而不覆盖。

    ``title`` 可选：作为 Title 样式段落写在文档开头；``markdown`` 为正文
    markdown 片段（标题/列表/粗体见 ``markdown_to_docx``）。父目录不存在时
    自动创建。
    """
    if os.path.exists(path):
        return {"success": False, "error": f"文件已存在: {path}"}
    parent = os.path.dirname(os.path.abspath(path))
    os.makedirs(parent, exist_ok=True)
    doc = Document()
    apply_academic_style(doc)
    if title:
        paragraph = doc.add_paragraph(style="Title")
        paragraph.add_run(title)
    markdown_to_docx(doc, markdown)
    doc.save(path)
    return {"success": True, "path": path}


def append_to_document(path: str, markdown: str) -> dict:
    """向已有 docx 追加 markdown 内容；文件不存在时报错。"""
    if not os.path.exists(path):
        return {"success": False, "error": f"文件不存在: {path}"}
    doc = Document(path)
    markdown_to_docx(doc, markdown)
    doc.save(path)
    return {"success": True, "path": path}


def reformat_document(path: str) -> dict:
    """重排已有 docx 为学术格式（委托 ``reformat_docx``，原地保存，幂等）。"""
    if not os.path.exists(path):
        return {"success": False, "error": f"文件不存在: {path}"}
    reformat_docx(path)
    return {"success": True, "path": path}
