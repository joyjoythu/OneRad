import json
import os
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from langchain_core.messages import AIMessage

from app.agent.nodes import _run_system_command, process_tool_calls
from app.agent.state import AgentState
from app.agent.tools import build_tools
from app.docx_style import (
    apply_academic_style,
    markdown_to_docx,
    reformat_docx,
    style_table,
)
from app.report import ReportAgent
from app.word_document import (
    append_to_document,
    create_document,
    reformat_document,
)

BLACK = RGBColor(0, 0, 0)


def _style_east(style):
    return style.element.rPr.rFonts.get(qn("w:eastAsia"))


def _run_east(run):
    return run._element.rPr.rFonts.get(qn("w:eastAsia"))


def _base_analysis_result():
    return {
        "success": True,
        "n_samples": 60,
        "selected_features": ["original_firstorder_Mean"],
        "model_results": {
            "intercept": 0.0,
            "coefficients": {"original_firstorder_Mean": 0.5},
            "odds_ratios": {"original_firstorder_Mean": 1.65},
            "ci_lower": {"original_firstorder_Mean": 1.0},
            "ci_upper": {"original_firstorder_Mean": 2.5},
            "p_values": {"original_firstorder_Mean": 0.01},
        },
        "metrics": {
            "auc": 0.85, "auc_ci": [0.74, 0.93],
            "accuracy": 0.8, "sensitivity": 0.82, "specificity": 0.78,
        },
    }


# ---------------------------------------------------------------------------
# apply_academic_style：样式级断言
# ---------------------------------------------------------------------------

def test_apply_academic_style_normal():
    doc = Document()
    apply_academic_style(doc)
    normal = doc.styles["Normal"]
    assert normal.font.name == "Times New Roman"
    assert _style_east(normal) == "宋体"
    assert normal.font.size == Pt(12)
    assert normal.font.color.rgb == BLACK
    assert normal.paragraph_format.line_spacing == 1.5


def test_apply_academic_style_title():
    doc = Document()
    apply_academic_style(doc)
    title = doc.styles["Title"]
    assert _style_east(title) == "黑体"
    assert title.font.size == Pt(16)
    assert title.font.bold is True
    assert title.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_apply_academic_style_headings_black_not_theme_blue():
    doc = Document()
    apply_academic_style(doc)
    h1 = doc.styles["Heading 1"]
    assert _style_east(h1) == "黑体"
    assert h1.font.size == Pt(14)
    assert h1.font.bold is True
    # 显式黑色：覆盖 python-docx 默认主题的蓝色 Heading
    assert h1.font.color.rgb == BLACK
    h2 = doc.styles["Heading 2"]
    assert _style_east(h2) == "黑体"
    assert h2.font.size == Pt(12)
    assert h2.font.bold is True
    assert h2.font.color.rgb == BLACK


# ---------------------------------------------------------------------------
# style_table：五号、表头加粗居中
# ---------------------------------------------------------------------------

def test_style_table_formats_header_and_body():
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "特征"
    table.rows[1].cells[0].text = "Mean"
    style_table(table)

    header_paragraph = table.rows[0].cells[0].paragraphs[0]
    header_run = header_paragraph.runs[0]
    assert header_run.font.size == Pt(10.5)
    assert header_run.font.bold is True
    assert header_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _run_east(header_run) == "宋体"

    body_run = table.rows[1].cells[0].paragraphs[0].runs[0]
    assert body_run.font.size == Pt(10.5)
    assert _run_east(body_run) == "宋体"


# ---------------------------------------------------------------------------
# markdown_to_docx：标题 / 粗体 / 列表渲染
# ---------------------------------------------------------------------------

def test_markdown_to_docx_renders_headings_bullets_and_bold():
    doc = Document()
    apply_academic_style(doc)
    markdown_to_docx(
        doc,
        "# 大标题\n\n## 一级小节\n\n### 二级小节\n\n"
        "- 要点 **一**\n\n正文 **加粗** 结尾\n",
    )
    styles = [p.style.name for p in doc.paragraphs]
    assert styles == ["Title", "Heading 1", "Heading 2",
                      "List Bullet", "Normal"]

    bullet = doc.paragraphs[3]
    assert "**" not in bullet.text
    bold_runs = [r for r in bullet.runs if r.bold]
    assert [r.text for r in bold_runs] == ["一"]

    body = doc.paragraphs[4]
    assert body.text == "正文 加粗 结尾"
    assert [r.text for r in body.runs if r.bold] == ["加粗"]


def test_markdown_to_docx_skips_blank_lines():
    doc = Document()
    markdown_to_docx(doc, "\n\n第一段\n\n\n第二段\n\n")
    assert [p.text for p in doc.paragraphs] == ["第一段", "第二段"]


# ---------------------------------------------------------------------------
# word_document：create / append / reformat
# ---------------------------------------------------------------------------

def test_create_document_with_title_and_refuse_overwrite(tmp_path):
    path = str(tmp_path / "notes.docx")
    result = create_document(path, "## 小节\n\n- 要点", title="文档标题")
    assert result["success"] is True
    doc = Document(path)
    assert doc.paragraphs[0].text == "文档标题"
    assert doc.paragraphs[0].style.name == "Title"
    # 新建文档已套用学术格式
    assert doc.styles["Normal"].font.size == Pt(12)

    again = create_document(path, "内容")
    assert again["success"] is False
    assert "已存在" in again["error"]


def test_append_to_document_requires_existing_file(tmp_path):
    missing = append_to_document(str(tmp_path / "missing.docx"), "内容")
    assert missing["success"] is False
    assert "不存在" in missing["error"]

    path = str(tmp_path / "doc.docx")
    create_document(path, "第一段")
    result = append_to_document(path, "- 追加要点")
    assert result["success"] is True
    doc = Document(path)
    assert doc.paragraphs[-1].text == "追加要点"
    assert doc.paragraphs[-1].style.name == "List Bullet"


def test_reformat_docx_normalizes_runs_in_place(tmp_path):
    # 构造一个格式混乱的文档：标题显式 Pt(18) 红色、正文显式 Pt(9)
    path = str(tmp_path / "messy.docx")
    doc = Document()
    heading = doc.add_heading(level=1)
    h_run = heading.add_run("混乱标题")
    h_run.font.size = Pt(18)
    h_run.font.color.rgb = RGBColor(0xFF, 0, 0)
    body = doc.add_paragraph()
    b_run = body.add_run("正文内容")
    b_run.font.size = Pt(9)
    bold_run = body.add_run("保留加粗")
    bold_run.bold = True
    doc.save(path)

    result = reformat_docx(path)
    assert result == path  # 原地保存，返回原路径
    assert not os.path.exists(str(tmp_path / "messy.bak.docx"))  # 不再生成备份

    fixed = Document(path)
    h_run = fixed.paragraphs[0].runs[0]
    assert h_run.font.size == Pt(14)  # Heading 1 四号
    assert h_run.font.color.rgb == BLACK
    assert _run_east(h_run) == "黑体"
    b_run, bold_run = fixed.paragraphs[1].runs
    assert b_run.font.size == Pt(12)
    assert _run_east(b_run) == "宋体"
    assert bold_run.bold is True  # 正文加粗不被抹掉

    # 幂等：再次重排结果一致
    assert reformat_docx(path) == path
    refixed = Document(path)
    assert refixed.paragraphs[0].runs[0].font.size == Pt(14)
    assert refixed.paragraphs[1].runs[0].font.size == Pt(12)


def test_reformat_document_delegates(tmp_path):
    path = str(tmp_path / "doc.docx")
    create_document(path, "内容")
    result = reformat_document(path)
    assert result["success"] is True
    assert result["path"] == path
    assert "backup" not in result
    assert not os.path.exists(str(tmp_path / "doc.bak.docx"))


# ---------------------------------------------------------------------------
# agent 工具：注册、路径安全、端到端执行
# ---------------------------------------------------------------------------

def test_word_tools_registered(tmp_path):
    tools = build_tools(str(tmp_path), MagicMock())
    for name in ("word_create", "word_append", "reformat_report"):
        assert name in tools
    result = json.loads(tools["word_create"].invoke(
        {"filename": "a.docx", "content_markdown": "x"}))
    assert result["_pending_tool"] == "word_create"


def test_word_tools_not_registered_in_readonly_mode(tmp_path):
    tools = build_tools(str(tmp_path), MagicMock(), readonly=True)
    for name in ("word_create", "word_append", "reformat_report"):
        assert name not in tools


def test_word_create_rejects_path_outside_project(tmp_path):
    result = _run_system_command(
        {"_pending_tool": "word_create",
         "args": {"filename": "../evil.docx", "content_markdown": "x"}},
        str(tmp_path),
    )
    assert "error" in result
    assert not os.path.exists(tmp_path.parent / "evil.docx")


def _tool_call_state(tmp_path, name, args):
    return AgentState(
        messages=[AIMessage(content="", tool_calls=[{
            "id": "tc-w1", "name": name, "args": args,
        }])],
        project_path=str(tmp_path),
        base_url="https://api.deepseek.com/v1",
        model="deepseek-v4-pro",
        api_key="test-key",
    )


def test_process_tool_calls_word_create_needs_confirmation(tmp_path):
    state = _tool_call_state(tmp_path, "word_create",
                             {"filename": "报告.docx", "content_markdown": "# 标题"})
    with patch("app.agent.nodes.ChatOpenAI") as mock_llm_cls:
        mock_llm_cls.return_value = MagicMock()
        updates = process_tool_calls(state)

    # 写操作：挂起等待用户确认，不直接落盘
    assert updates["interrupt_type"] == "system_command"
    pending = updates["pending_command"]
    assert pending["_pending_tool"] == "word_create"
    assert not (tmp_path / "报告.docx").exists()

    # 用户确认后执行：文档创建在项目内
    result = _run_system_command(pending, str(tmp_path))
    assert result["result"]["success"] is True
    doc = Document(str(tmp_path / "报告.docx"))
    assert doc.paragraphs[0].style.name == "Title"


def test_process_tool_calls_reformat_report_runs_immediately(tmp_path):
    out_dir = tmp_path / "radiomics_analysis"
    out_dir.mkdir()
    (out_dir / "analysis_result.json").write_text(
        json.dumps(_base_analysis_result(), ensure_ascii=False),
        encoding="utf-8")
    report = ReportAgent().run(
        analysis_result=_base_analysis_result(),
        output_dir=str(out_dir),
        modality="CT", n_features=107, covariates=[])
    assert report["success"] is True

    state = _tool_call_state(tmp_path, "reformat_report", {})
    with patch("app.agent.nodes.ChatOpenAI") as mock_llm_cls:
        mock_llm_cls.return_value = MagicMock()
        updates = process_tool_calls(state)

    # 免确认：不产生中断，直接返回执行结果
    assert updates["interrupt_type"] is None
    content = json.loads(updates["messages"][0].content)
    assert content["success"] is True
    assert content["report_path"].endswith("AutoRadiomics_Report.docx")
    assert "backup" not in content
    assert not (out_dir / "AutoRadiomics_Report.bak.docx").exists()


def test_process_tool_calls_reformat_report_without_analysis_dir(tmp_path):
    state = _tool_call_state(tmp_path, "reformat_report", {})
    with patch("app.agent.nodes.ChatOpenAI") as mock_llm_cls:
        mock_llm_cls.return_value = MagicMock()
        updates = process_tool_calls(state)
    content = json.loads(updates["messages"][0].content)
    assert content["success"] is False
    assert "analysis_result.json" in content["error"]


# ---------------------------------------------------------------------------
# ReportAgent 集成：解读小节的粗体保留（markdown_to_docx 行为变化）
# ---------------------------------------------------------------------------

def test_report_interpretation_preserves_bold(tmp_path):
    interpretation = {
        "performance": "AUC 为 **0.85**，判别能力较好。",
        "features": "",
        "shap": "",
    }
    result = ReportAgent().run(
        analysis_result=_base_analysis_result(),
        output_dir=str(tmp_path),
        modality="CT", n_features=107, covariates=[],
        interpretation=interpretation)
    assert result["success"] is True
    doc = Document(result["report_path"])
    paragraph = next(p for p in doc.paragraphs if "判别能力" in p.text)
    assert "**" not in paragraph.text
    assert [r.text for r in paragraph.runs if r.bold] == ["0.85"]
