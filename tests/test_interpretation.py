import json
import os
from unittest.mock import MagicMock, patch

import numpy as np
import pandas as pd
import pytest
from docx import Document
from langchain_core.messages import AIMessage

from app.agent.nodes import process_tool_calls
from app.agent.state import AgentState
from app.agent.tools import build_tools
from app.interpretation import (
    SECTION_MARKERS,
    _parse_sections,
    apply_to_reports,
    build_summary,
    interpret,
    parse_feature_name,
)
from app.radiomics_analysis import (
    _render_markdown_report,
    run_radiomics_cv_analysis,
)
from app.report import ReportAgent


# ---------------------------------------------------------------------------
# 公共构造
# ---------------------------------------------------------------------------

def _make_run_inputs(tmp_path, n=60, seed=42):
    ids = [f"P{i:03d}" for i in range(n)]
    rng = np.random.RandomState(seed)
    label = np.array([i % 2 for i in range(n)])
    feat = pd.DataFrame({"patient_id": ids})
    for j in range(6):
        feat[f"original_sig_{j}"] = rng.randn(n) + label * 1.5
    feat_path = tmp_path / "features.csv"
    feat.to_csv(feat_path, index=False)
    clin_path = tmp_path / "clinical.csv"
    pd.DataFrame({"patient_id": ids, "Label": label}).to_csv(clin_path, index=False)
    return str(feat_path), str(clin_path)


def _base_analysis_result():
    return {
        "success": True,
        "n_samples": 60,
        "selected_features": ["wavelet-LLH_glcm_Contrast",
                              "original_firstorder_Mean"],
        "model_results": {
            "intercept": 0.1,
            "coefficients": {"wavelet-LLH_glcm_Contrast": 0.8,
                             "original_firstorder_Mean": -0.5},
            "odds_ratios": {"wavelet-LLH_glcm_Contrast": 2.23,
                            "original_firstorder_Mean": 0.61},
            "ci_lower": {"wavelet-LLH_glcm_Contrast": 1.2,
                         "original_firstorder_Mean": 0.3},
            "ci_upper": {"wavelet-LLH_glcm_Contrast": 4.1,
                         "original_firstorder_Mean": 1.2},
            "p_values": {"wavelet-LLH_glcm_Contrast": 0.01,
                         "original_firstorder_Mean": 0.14},
        },
        "metrics": {
            "auc": 0.85, "auc_ci": [0.74, 0.93],
            "accuracy": 0.8, "sensitivity": 0.82, "specificity": 0.78,
            "ppv": 0.79, "npv": 0.81, "f1": 0.8,
            "threshold": 0.5, "confusion_matrix": [[24, 6], [6, 24]],
        },
        "cv_metrics": {
            "folds": [
                {"fold": i + 1, "auc": 0.8, "accuracy": 0.8, "sensitivity": 0.8,
                 "specificity": 0.8, "ppv": 0.8, "npv": 0.8, "f1": 0.8,
                 "threshold": 0.5}
                for i in range(5)
            ],
            "mean": {"auc": 0.8, "accuracy": 0.8, "sensitivity": 0.8,
                     "specificity": 0.8, "ppv": 0.8, "npv": 0.8, "f1": 0.8,
                     "threshold": 0.5},
            "std": {"auc": 0.02, "accuracy": 0.02, "sensitivity": 0.02,
                    "specificity": 0.02, "ppv": 0.02, "npv": 0.02, "f1": 0.02,
                    "threshold": 0.01},
        },
        "oof_probabilities": [0.2, 0.7] * 30,
    }


def _fake_interpretation():
    return {
        "performance": "AUC 为 0.85，判别能力较好。",
        "features": "- wavelet-LLH_glcm_Contrast：系数为正，OR>1。",
        "shap": "wavelet-LLH_glcm_Contrast 的 mean|SHAP| 最大。",
    }


def _marked_llm_text():
    interp = _fake_interpretation()
    return (
        f"{SECTION_MARKERS['performance']}\n{interp['performance']}\n"
        f"{SECTION_MARKERS['features']}\n{interp['features']}\n"
        f"{SECTION_MARKERS['shap']}\n{interp['shap']}\n"
        "局限性：相关性非因果，需外部验证。"
    )


def _mock_llm_client(text=None):
    client = MagicMock()
    client.call.return_value = text if text is not None else _marked_llm_text()
    return client


def _write_shap_csvs(output_dir):
    """合成两折 SHAP CSV：Contrast 两折均在且重要性最高，Mean 仅 fold1。"""
    shap_dir = os.path.join(output_dir, "shap")
    os.makedirs(shap_dir, exist_ok=True)
    pd.DataFrame({
        "patient_id": ["P1", "P2"],
        "wavelet-LLH_glcm_Contrast": [0.5, 0.7],   # mean_abs 0.6, 方向正
        "original_firstorder_Mean": [0.1, -0.1],   # mean_abs 0.1, 方向≈0
    }).to_csv(os.path.join(shap_dir, "shap_values_fold1.csv"), index=False)
    pd.DataFrame({
        "patient_id": ["P3", "P4"],
        "wavelet-LLH_glcm_Contrast": [0.8, 1.0],   # mean_abs 0.9, 方向正
    }).to_csv(os.path.join(shap_dir, "shap_values_fold2.csv"), index=False)


# ---------------------------------------------------------------------------
# analysis_result.json 落盘
# ---------------------------------------------------------------------------

def test_analysis_result_json_written_and_loadable(tmp_path):
    feat, clin = _make_run_inputs(tmp_path)
    out_dir = str(tmp_path / "analysis_out")
    result = run_radiomics_cv_analysis(
        feature_csv=feat, clinical=clin, output_dir=out_dir,
        id_col="patient_id", label_col="Label")
    assert result["success"] is True

    json_path = result["outputs"]["analysis_result_json"]
    assert json_path == os.path.join(out_dir, "analysis_result.json")
    with open(json_path, encoding="utf-8") as f:
        payload = json.load(f)
    assert set(payload) == {"metrics", "cv_metrics", "model_results",
                            "selected_features", "n_samples",
                            "oof_probabilities"}
    assert payload["n_samples"] == 60
    assert payload["metrics"]["auc"] > 0
    assert len(payload["oof_probabilities"]) == 60


# ---------------------------------------------------------------------------
# 特征名解析与 build_summary 聚合
# ---------------------------------------------------------------------------

def test_parse_feature_name():
    assert parse_feature_name("wavelet-LLH_glcm_Contrast") == {
        "filter": "wavelet-LLH", "feature_class": "glcm",
        "feature_name": "Contrast"}
    assert parse_feature_name("original_firstorder_Mean") == {
        "filter": "original", "feature_class": "firstorder",
        "feature_name": "Mean"}
    assert parse_feature_name("shape_Volume") == {
        "filter": "", "feature_class": "shape", "feature_name": "Volume"}
    assert parse_feature_name("plain") == {
        "filter": "", "feature_class": "", "feature_name": "plain"}


def test_build_summary_aggregates_shap(tmp_path):
    out_dir = str(tmp_path)
    _write_shap_csvs(out_dir)
    summary = build_summary(_base_analysis_result(), out_dir)

    assert summary["n_samples"] == 60
    assert summary["n_positive"] == 30
    assert summary["n_negative"] == 30
    assert summary["metrics"]["auc"] == 0.85

    feats = {f["feature"]: f for f in summary["features"]}
    assert feats["wavelet-LLH_glcm_Contrast"]["filter"] == "wavelet-LLH"
    assert feats["wavelet-LLH_glcm_Contrast"]["feature_class"] == "glcm"
    assert feats["wavelet-LLH_glcm_Contrast"]["coefficient_direction"] == "positive"
    assert feats["original_firstorder_Mean"]["coefficient_direction"] == "negative"

    shap = summary["shap"]
    assert shap["n_folds"] == 2
    top = shap["top_features"]
    # Contrast 跨折 mean|SHAP| = (0.6 + 0.9) / 2 = 0.75，排第一
    assert top[0]["feature"] == "wavelet-LLH_glcm_Contrast"
    assert top[0]["mean_abs_shap"] == pytest.approx(0.75)
    assert top[0]["fold_coverage"] == 2
    assert top[0]["shap_direction"] == "positive"
    # 系数方向（正）与 SHAP 方向（正）一致
    assert top[0]["direction_consistent"] is True
    # Mean 只在 fold1 出现，覆盖次数 1；系数为负、SHAP 均值≈0 → 无法判断一致性
    mean_entry = [t for t in top
                  if t["feature"] == "original_firstorder_Mean"][0]
    assert mean_entry["fold_coverage"] == 1
    assert mean_entry["mean_abs_shap"] == pytest.approx(0.1)
    assert mean_entry["direction_consistent"] is None


def test_build_summary_without_shap_dir(tmp_path):
    summary = build_summary(_base_analysis_result(), str(tmp_path))
    assert summary["shap"]["n_folds"] == 0
    assert summary["shap"]["top_features"] == []


# ---------------------------------------------------------------------------
# interpret：mock LLMClient，分隔标记解析
# ---------------------------------------------------------------------------

def test_interpret_parses_marked_sections():
    client = _mock_llm_client()
    sections = interpret({"metrics": {}}, client)
    assert sections == _fake_interpretation() | {
        "shap": _fake_interpretation()["shap"] + "\n局限性：相关性非因果，需外部验证。"}
    client.call.assert_called_once()
    system, user = client.call.call_args.args
    assert "【模型性能解读】" in system
    assert "数值摘要" in user


def test_interpret_raises_on_llm_exception():
    client = MagicMock()
    client.call.side_effect = RuntimeError("LLMClient 未配置 API key")
    with pytest.raises(RuntimeError):
        interpret({}, client)


def test_interpret_raises_on_missing_marker():
    client = _mock_llm_client("没有任何分隔标记的自由文本")
    with pytest.raises(ValueError, match="分隔标记"):
        interpret({}, client)


def test_parse_sections_rejects_reordered_or_empty():
    m = SECTION_MARKERS
    with pytest.raises(ValueError):
        _parse_sections(f"{m['features']}x\n{m['performance']}y\n{m['shap']}z")
    with pytest.raises(ValueError, match="为空"):
        _parse_sections(f"{m['performance']}\n{m['features']}x\n{m['shap']}z")


# ---------------------------------------------------------------------------
# 报告注入：md / docx / 幂等
# ---------------------------------------------------------------------------

def test_markdown_report_injects_interpretation(tmp_path):
    path = _render_markdown_report(
        _base_analysis_result(), outputs={}, n_matched=60, covariates=[],
        output_dir=str(tmp_path), n_splits=5,
        interpretation=_fake_interpretation())
    md = open(path, encoding="utf-8").read()
    assert "## 6. 结果解读" in md
    assert "### 模型性能解读" in md
    assert "### 特征意义解读" in md
    assert "### SHAP 可解释性解读" in md
    assert "AUC 为 0.85" in md


def test_markdown_report_without_interpretation_unchanged(tmp_path):
    path = _render_markdown_report(
        _base_analysis_result(), outputs={}, n_matched=60, covariates=[],
        output_dir=str(tmp_path), n_splits=5)
    md = open(path, encoding="utf-8").read()
    assert "结果解读" not in md


def test_docx_report_injects_interpretation(tmp_path):
    result = ReportAgent().run(
        analysis_result=_base_analysis_result(),
        output_dir=str(tmp_path),
        modality="auto", n_features=107, covariates=[],
        interpretation=_fake_interpretation())
    assert result["success"] is True
    doc = Document(result["report_path"])
    texts = [p.text for p in doc.paragraphs]
    assert "结果解读" in texts
    assert "模型性能解读" in texts
    assert "SHAP 可解释性解读" in texts
    assert any("AUC 为 0.85" in t for t in texts)


def test_docx_report_without_interpretation_unchanged(tmp_path):
    result = ReportAgent().run(
        analysis_result=_base_analysis_result(),
        output_dir=str(tmp_path),
        modality="auto", n_features=107, covariates=[])
    assert result["success"] is True
    doc = Document(result["report_path"])
    assert "结果解读" not in [p.text for p in doc.paragraphs]


def test_apply_to_reports_idempotent(tmp_path):
    out_dir = str(tmp_path)
    interp = _fake_interpretation()
    for _ in range(2):
        reports = apply_to_reports(_base_analysis_result(), out_dir, interp)
    md = open(reports["report_md"], encoding="utf-8").read()
    assert md.count("## 6. 结果解读") == 1
    assert md.count("### 模型性能解读") == 1
    doc = Document(reports["report_docx"])
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert headings.count("结果解读") == 1


# ---------------------------------------------------------------------------
# agent 工具：注册与免确认执行
# ---------------------------------------------------------------------------

def test_interpret_tool_registered(tmp_path):
    tools = build_tools(str(tmp_path), MagicMock())
    assert "interpret_analysis_results" in tools
    result = json.loads(tools["interpret_analysis_results"].invoke({}))
    assert result["_pending_tool"] == "interpret_analysis_results"


def test_interpret_tool_not_registered_in_readonly_mode(tmp_path):
    tools = build_tools(str(tmp_path), MagicMock(), readonly=True)
    assert "interpret_analysis_results" not in tools


def _tool_call_state(tmp_path):
    return AgentState(
        messages=[AIMessage(content="", tool_calls=[{
            "id": "tc-i1",
            "name": "interpret_analysis_results",
            "args": {},
        }])],
        project_path=str(tmp_path),
        base_url="https://api.deepseek.com/v1",
        model="deepseek-v4-pro",
        api_key="test-key",
    )


def _run_process(state, llm_client):
    with patch("app.agent.nodes.ChatOpenAI") as mock_llm_cls, \
         patch("app.llm.LLMClient", return_value=llm_client):
        mock_llm_cls.return_value = MagicMock()
        return process_tool_calls(state)


def _write_analysis_dir(tmp_path, with_params=True):
    out_dir = tmp_path / "radiomics_analysis"
    out_dir.mkdir()
    payload = _base_analysis_result()
    payload.pop("success")
    payload.pop("oof_probabilities")
    (out_dir / "analysis_result.json").write_text(
        json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    if with_params:
        (out_dir / "analysis_params.json").write_text(json.dumps({
            "feature_csv": "features.csv", "clinical": "clinical.csv",
            "output_dir": "radiomics_analysis", "project_root": "..",
            "id_col": "patient_id", "label_col": "Label", "covariates": [],
            "max_lasso_features": 100, "n_splits": 5, "random_state": 42,
        }, ensure_ascii=False), encoding="utf-8")
    return out_dir


def test_process_tool_calls_executes_interpretation_immediately(tmp_path):
    out_dir = _write_analysis_dir(tmp_path)
    _write_shap_csvs(str(out_dir))
    updates = _run_process(_tool_call_state(tmp_path), _mock_llm_client())

    # 免确认：不产生中断，直接返回执行结果 ToolMessage
    assert updates["interrupt_type"] is None
    assert len(updates["messages"]) == 1
    content = json.loads(updates["messages"][0].content)
    assert content["success"] is True
    outputs = content["outputs"]
    assert outputs["interpretation"].endswith("interpretation.md")

    md = open(outputs["report_md"], encoding="utf-8").read()
    assert "## 6. 结果解读" in md
    doc = Document(outputs["report_docx"])
    assert "结果解读" in [p.text for p in doc.paragraphs]
    interp_md = open(outputs["interpretation"], encoding="utf-8").read()
    assert "模型性能解读" in interp_md

    # 重复执行幂等：报告不叠加小节
    updates = _run_process(_tool_call_state(tmp_path), _mock_llm_client())
    content = json.loads(updates["messages"][0].content)
    md = open(content["outputs"]["report_md"], encoding="utf-8").read()
    assert md.count("## 6. 结果解读") == 1


def test_process_tool_calls_missing_analysis_result_returns_error(tmp_path):
    updates = _run_process(_tool_call_state(tmp_path), _mock_llm_client())
    assert updates["interrupt_type"] is None
    content = json.loads(updates["messages"][0].content)
    assert content["success"] is False
    assert "analysis_result.json" in content["error"]
    assert "重新运行分析" in content["error"]


def test_process_tool_calls_llm_failure_keeps_base_report(tmp_path):
    out_dir = _write_analysis_dir(tmp_path)
    base_report = out_dir / "report.md"
    base_report.write_text("# 基础报告\n", encoding="utf-8")

    client = MagicMock()
    client.call.side_effect = RuntimeError("LLMClient 未配置 API key")
    updates = _run_process(_tool_call_state(tmp_path), client)

    content = json.loads(updates["messages"][0].content)
    assert content["success"] is False
    assert "结果解读失败" in content["error"]
    assert "不受影响" in content["error"]
    # 基础报告未被改动
    assert base_report.read_text(encoding="utf-8") == "# 基础报告\n"
    assert not (out_dir / "interpretation.md").exists()
