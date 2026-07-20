from pathlib import Path

import pytest
from docx import Document

from app.report import ReportAgent


def _base_analysis_result():
    return {
        "success": True,
        "task_type": "binary_classification",
        "selected_features": ["original_firstorder_Mean"],
        "model_results": {
            "intercept": 0.0,
            "coefficients": {"original_firstorder_Mean": 0.5},
            "odds_ratios": {"original_firstorder_Mean": 1.65},
            "ci_lower": {"original_firstorder_Mean": 1.0},
            "ci_upper": {"original_firstorder_Mean": 2.5},
            "p_values": {"original_firstorder_Mean": 0.01},
        },
        "metrics": {
            "auc": 0.85,
            "auc_ci": [0.78, 0.91],
            "accuracy": 0.80,
            "sensitivity": 0.82,
            "specificity": 0.78,
            "threshold": 0.5,
            "confusion_matrix": [[40, 10], [8, 42]],
        },
        "n_samples": 100,
    }


def test_report_generation(tmp_path):
    analysis_result = _base_analysis_result()
    agent = ReportAgent()
    result = agent.run(
        analysis_result=analysis_result,
        output_dir=str(tmp_path),
        modality="CT",
        n_features=107,
        covariates=[],
    )
    assert result["success"] is True
    assert Path(result["report_path"]).exists()
    assert result["skipped_plots"] == []


def test_report_generation_failed_upstream(tmp_path):
    analysis_result = {"success": False, "message": "upstream analysis failed"}
    result = ReportAgent().run(
        analysis_result=analysis_result,
        output_dir=str(tmp_path),
        modality="CT",
        n_features=107,
        covariates=[],
    )
    assert result["success"] is False
    assert "上游分析失败" in result["message"]
    assert "upstream analysis failed" in result["message"]


@pytest.mark.parametrize(
    "missing_key",
    [
        "n_samples",
        "selected_features",
        "model_results",
        "metrics",
    ],
)
def test_report_generation_missing_top_level_key(tmp_path, missing_key):
    analysis_result = _base_analysis_result()
    del analysis_result[missing_key]
    result = ReportAgent().run(
        analysis_result=analysis_result,
        output_dir=str(tmp_path),
        modality="CT",
        n_features=107,
        covariates=[],
    )
    assert result["success"] is False
    assert "analysis_result 缺少字段" in result["message"]
    assert missing_key in result["message"]


@pytest.mark.parametrize(
    "nested_key",
    [
        "coefficients",
        "odds_ratios",
        "ci_lower",
        "ci_upper",
        "p_values",
    ],
)
def test_report_generation_missing_model_result_key(tmp_path, nested_key):
    analysis_result = _base_analysis_result()
    del analysis_result["model_results"][nested_key]
    result = ReportAgent().run(
        analysis_result=analysis_result,
        output_dir=str(tmp_path),
        modality="CT",
        n_features=107,
        covariates=[],
    )
    assert result["success"] is False
    assert "analysis_result 缺少字段" in result["message"]
    assert f"model_results.{nested_key}" in result["message"]


@pytest.mark.parametrize(
    "nested_key",
    [
        "auc",
        "auc_ci",
        "accuracy",
        "sensitivity",
        "specificity",
    ],
)
def test_report_generation_missing_metric_key(tmp_path, nested_key):
    analysis_result = _base_analysis_result()
    del analysis_result["metrics"][nested_key]
    result = ReportAgent().run(
        analysis_result=analysis_result,
        output_dir=str(tmp_path),
        modality="CT",
        n_features=107,
        covariates=[],
    )
    assert result["success"] is False
    assert "analysis_result 缺少字段" in result["message"]
    assert f"metrics.{nested_key}" in result["message"]


class MockLLMClient:
    def call(self, system, prompt, **kwargs):
        return "Polished methodology paragraph containing all exact numbers."


def test_report_generation_llm_polish(tmp_path):
    analysis_result = _base_analysis_result()
    result = ReportAgent().run(
        analysis_result=analysis_result,
        output_dir=str(tmp_path),
        modality="CT",
        n_features=107,
        covariates=["age", "gender"],
        llm_client=MockLLMClient(),
    )
    assert result["success"] is True
    doc = Document(result["report_path"])
    paragraphs = [p.text for p in doc.paragraphs]
    assert any("Polished methodology paragraph" in text for text in paragraphs)


def test_report_generation_missing_plot_is_skipped(tmp_path):
    analysis_result = _base_analysis_result()
    missing_plot = str(tmp_path / "missing_plot.png")
    result = ReportAgent().run(
        analysis_result=analysis_result,
        output_dir=str(tmp_path),
        modality="CT",
        n_features=107,
        covariates=[],
        plot_paths=[missing_plot],
    )
    assert result["success"] is True
    assert missing_plot in result["skipped_plots"]


def test_report_generation_empty_selected_features(tmp_path):
    analysis_result = _base_analysis_result()
    analysis_result["selected_features"] = []
    result = ReportAgent().run(
        analysis_result=analysis_result,
        output_dir=str(tmp_path),
        modality="CT",
        n_features=107,
        covariates=[],
    )
    assert result["success"] is True
    assert Path(result["report_path"]).exists()


def _docx_all_text(path):
    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts)


def test_report_generation_with_cv_metrics_and_small_p(tmp_path):
    """cv_metrics 存在时渲染逐折表格；p<0.001 归并显示。"""
    analysis_result = _base_analysis_result()
    analysis_result["model_results"]["p_values"]["original_firstorder_Mean"] = 0.0004
    analysis_result["metrics"].update({"ppv": 0.81, "npv": 0.79, "f1": 0.80})
    metric_keys = ("auc", "accuracy", "sensitivity", "specificity",
                   "ppv", "npv", "f1", "threshold")
    analysis_result["cv_metrics"] = {
        "folds": [
            dict({"fold": i + 1}, **{k: 0.8 for k in metric_keys})
            for i in range(5)
        ],
        "mean": {k: 0.8 for k in metric_keys},
        "std": {k: 0.02 for k in metric_keys},
    }
    result = ReportAgent().run(
        analysis_result=analysis_result,
        output_dir=str(tmp_path),
        modality="CT",
        n_features=107,
        covariates=[],
    )
    assert result["success"] is True
    text = _docx_all_text(result["report_path"])
    assert "<0.001" in text
    assert "0.0004" not in text
    assert "Mean±SD" in text
    assert "0.800±0.020" in text
    assert "PPV = 0.810" in text


def test_report_generation_without_cv_metrics_still_works(tmp_path):
    """缺少 cv_metrics 的旧输入仍能生成报告（无逐折表）。"""
    result = ReportAgent().run(
        analysis_result=_base_analysis_result(),
        output_dir=str(tmp_path),
        modality="CT",
        n_features=107,
        covariates=[],
    )
    assert result["success"] is True
    text = _docx_all_text(result["report_path"])
    assert "Mean±SD" not in text
